"""Word 解析器——使用 python-docx

内存优化：
- try/finally 确保 Document 对象释放
- 表格存储用 list[list[str]] 替代 list[list[PhysicalCell]]（与 PdfParser 对齐）
- 单表行数限制（防止超大 Word 表格撑爆内存）
- raw_text 与 paragraphs 共享数据（不重复存储）
"""

from pathlib import Path

# python-docx 改为懒导入——在 parse() 中导入，节省模块导入内存

from .base import BaseParser, ParsedDocument


# 单表最大行数（防超大表格撑爆内存）
_MAX_TABLE_ROWS = 500


class WordParser(BaseParser):
    """用 python-docx 提取段落、表格（紧凑存储）"""

    def parse(self, file_path: str) -> ParsedDocument:
        paragraphs: list[str] = []
        # 紧凑存储 list[list[list[str]]]，与 PdfParser 对齐
        tables: list[list[list[str]]] = []
        sheets_info: list[dict] = []

        # 懒导入 python-docx（节省模块导入内存）
        from docx import Document
        doc = Document(file_path)
        try:
            # 提取段落
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            # 提取表格（紧凑存储，限制行数）
            for table in doc.tables:
                current_table: list[list[str]] = []
                for ri, row in enumerate(table.rows):
                    if ri >= _MAX_TABLE_ROWS:
                        current_table.append([f"...（已截断，共 {len(table.rows)} 行）"])
                        break
                    current_row = [cell.text.strip() for cell in row.cells]
                    if current_row:
                        current_table.append(current_row)
                if current_table:
                    tables.append(current_table)
                    sheets_info.append({
                        "name": f"table_{len(tables)}",
                        "rows": len(current_table),
                        "cols": len(current_table[0]) if current_table else 0,
                    })
        finally:
            # python-docx 的 Document 没有显式 close 方法，
            # 但通过 del 解除引用让 GC 尽早回收
            try:
                del doc
            except Exception:
                pass

        # raw_text 由 paragraphs + 表格文本拼接（不重复存储段落）
        parts = list(paragraphs)
        for table in tables:
            for row in table:
                parts.append(" | ".join(row))
        raw_text = "\n".join(parts)
        # 释放临时列表
        del parts

        return ParsedDocument(
            raw_text=raw_text,
            tables=tables,
            paragraphs=paragraphs,
            metadata={
                "filename": Path(file_path).name,
                "paragraphs_count": len(paragraphs),
                "tables_count": len(tables),
                "tables_info": sheets_info,
            }
        )
