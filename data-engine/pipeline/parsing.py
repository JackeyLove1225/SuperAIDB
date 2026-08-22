"""解析层：文件 → 文本流（PDF/Excel/Word 统一流单元契约）

流单元契约 iter[(unit_no, text, table_count)]：PDF unit=页、Excel unit=行块、
Word unit=段批/表批。下游（tier-1 向量分流/tier-2 AI 提取）对来源格式零感知。
"""
from core.logger import get_logger
from pathlib import Path

import os

from pipeline.constants import TIER1_EXCEL_ROWS, TIER1_DOCX_PARAS

logger = get_logger(__name__)

def excel_to_text_stream(file_path: str, sheet_start: int = 0, sheet_limit: int = None):
    """流式 Excel 提取——按流单元 yield（流单元 = ≤TIER1_EXCEL_ROWS 行块）

    与 PDF 路线对齐：通用单元格提取，结构化交给 AI FC。
    行数超过 TIER1_EXCEL_ROWS 的 sheet 按行切块，跨块携带首行表头。

    Args:
        file_path: Excel 文件路径
        sheet_start: 起始流单元序号（0-indexed，按全文件流单元计数）
        sheet_limit: 最多处理的流单元数（None=全部）
    Yields:
        (unit_no, unit_text, 1)  — tc=1：Excel 流单元即表格单元，
        走 tier-1 暂存 → AI 提取 → 剥离后叙述文本进向量库的路径
    """
    import openpyxl
    from core.parser.excel_parser import ExcelParser

    wb = openpyxl.load_workbook(file_path, data_only=True)
    parser = ExcelParser()
    start = max(0, sheet_start)
    unit_idx = 0   # 全文件流单元序号（跨 sheet 连续计数）
    emitted = 0
    try:
        for ws in wb.worksheets:
            merge_map = parser._build_merge_map(ws)
            for text in parser.iter_sheet_units(ws, merge_map, TIER1_EXCEL_ROWS):
                if unit_idx >= start and (sheet_limit is None or emitted < sheet_limit):
                    emitted += 1
                    yield unit_idx + 1, text, 1
                unit_idx += 1
    finally:
        wb.close()


def pdf_to_text_stream(pdf_path: str, page_start: int = 0, page_limit: int = None,
                        extract_tables: bool = True):
    """流式提取 PDF 文本和表格

    迁移说明（2026-07-19）：
    旧实现走 core.pdf_to_text.pdf_to_text_stream（pdfplumber + PyMuPDF 双实现，76秒/476页）
    新实现走 core.parser.pdf_parser.PdfParser.parse_stream（纯 PyMuPDF，46秒/476页，-40%）
    - 输出格式不变：(page_num, page_text, table_count)
    - page_text 从带坐标的详细格式改为纯文本（向量入库和 AI 分析都不需要坐标噪音）
    - 表格检测从启发式（≥3行≥2列）改为 find_tables() 真实检测
    """
    from core.parser.pdf_parser import PdfParser
    parser = PdfParser(extract_tables=extract_tables)
    return parser.parse_stream(pdf_path, page_start=page_start, page_limit=page_limit)


def docx_to_text_stream(docx_path: str, page_start: int = 0, page_limit: int = None):
    """流式提取 Word 文档文本——按流单元 yield（流单元 = 段批或单表批）

    Word 文档没有页码概念，这里将"段落+表格"按顺序组织：
    - 段落连续内容合并为一批（每批最多 TIER1_DOCX_PARAS 段），tc=0 纯文字单元
    - 每个表格单独作为一批，tc=1 表格单元（走 AI 判定：业务表→关系库，
      说明性零散表格→剥离后原文进向量库）
    - page_start/page_limit 按流单元序号切片（0-indexed，None=全部）

    已知限制（P2）：表格批排在段落批之后，脱离原文位置（python-docx 段落/表格
    分离枚举所致），demo 后排期做位置回填。

    Yields:
        (unit_no, text, tc)  — tc: 0=纯文字单元，1=表格单元
    """
    from docx import Document

    doc = Document(docx_path)
    batches = []  # [(text, tc), ...]

    # 收集段落文本（非空段落）
    paragraphs_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # 合并连续段落为一批（每批最多 TIER1_DOCX_PARAS 段，模拟"一页"内容量）
    for i in range(0, len(paragraphs_text), TIER1_DOCX_PARAS):
        chunk = "\n".join(paragraphs_text[i:i + TIER1_DOCX_PARAS])
        if chunk:
            batches.append((chunk, 0))

    # 收集表格文本（每个表格作为一批，格式与 Excel 一致：| 分隔）
    for table in doc.tables:
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows_text.append(" | ".join(cells))
        if rows_text:
            batches.append(("\n".join(rows_text), 1))

    # 应用 page_start 和 page_limit
    start = max(0, page_start)
    end = len(batches) if not page_limit else min(len(batches), start + page_limit)

    for idx in range(start, end):
        text, tc = batches[idx]
        yield idx + 1, text, tc




# 文件大小限制（MB）
MAX_FILE_SIZE_MB = 50


def _check_file_size(file_path: str) -> None:
    """检查文件大小是否超过限制

    Raises:
        ValueError: 文件超过 50MB 限制
    """
    import os
    size_mb = os.path.getsize(file_path) / 1024 / 1024
    if size_mb > MAX_FILE_SIZE_MB:
        raise ValueError(
            f"文件过大：{size_mb:.1f}MB，超过 {MAX_FILE_SIZE_MB}MB 限制。"
            f"请拆分文件后重新上传。"
        )


# ── 分批处理 ──

# ── 图片/扫描件 OCR 路径（tier-1 平级：只产文字，上层零感知）──

OCR_MIN_TEXT_LEN = 20  # PDF 页文本少于此长度视为无文字层（扫描页），走 OCR 回退


def image_to_text_stream(image_path: str, page_start: int = 0, page_limit: int = None):
    """图片文件（png/jpg/jpeg/bmp/webp）→ OCR 文字流（单流单元契约）

    与 pdf/excel/word 平级：产出 (1, text, tc)，tc 按内容粗判（含数字/表格线=1）。
    OCR 未配置/失败时如实抛错（不静默产出空文本冒充成功）。
    """
    from pipeline.ocr import ocr_image_to_markdown, is_available
    if not is_available():
        raise ValueError("图片识别需要 OCR，但 config/.env 未配置 OCR_API_TOKEN。"
                         "文字版 PDF/Word/Excel 不受影响，可直接处理。")
    text = ocr_image_to_markdown(image_path)
    tc = 1 if any(ch.isdigit() for ch in text) else 0
    yield 1, text, tc


def pdf_to_text_stream_with_ocr(pdf_path: str, page_start: int = 0, page_limit: int = None,
                                extract_tables: bool = True):
    """PDF 文本流（扫描页 OCR 回退版）

    与 pdf_to_text_stream 同契约；某页提取文本 < OCR_MIN_TEXT_LEN 时判定为
    扫描页/无文字层，PyMuPDF 渲染为图片后走 PaddleOCR 取文字。
    OCR 未配置时扫描页按原文（通常为空）如实上报，不硬解。
    """
    from pipeline.ocr import ocr_image_to_markdown, is_available
    import fitz
    import tempfile

    base_stream = pdf_to_text_stream(pdf_path, page_start=page_start,
                                     page_limit=page_limit, extract_tables=extract_tables)
    doc = None
    for unit_no, text, tc in base_stream:
        if len((text or "").strip()) >= OCR_MIN_TEXT_LEN:
            yield unit_no, text, tc
            continue
        # 扫描页回退：渲染为图片 → OCR
        if not is_available():
            logger.warning("第 %s 页文本过短（疑似扫描件）且 OCR 未配置，按空文本上报", unit_no)
            yield unit_no, text, tc
            continue
        try:
            if doc is None:
                doc = fitz.open(pdf_path)
            page = doc[unit_no - 1]
            pix = page.get_pixmap(dpi=200)
            # Windows 下 NamedTemporaryFile 句柄未关会锁文件——先命名并立即关闭
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
            tmp_path = tmp.name
            tmp.close()
            try:
                pix.save(tmp_path)
                ocr_text = ocr_image_to_markdown(tmp_path)
            finally:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
            if ocr_text.strip():
                logger.info("第 %s 页扫描件 OCR 成功（%d 字）", unit_no, len(ocr_text))
                yield unit_no, ocr_text, 1 if tc or any(c.isdigit() for c in ocr_text) else 0
            else:
                yield unit_no, text, tc
        except Exception as e:
            logger.warning("第 %s 页 OCR 回退失败（按原文上报）: %s", unit_no, str(e)[:80])
            yield unit_no, text, tc
    if doc is not None:
        doc.close()
