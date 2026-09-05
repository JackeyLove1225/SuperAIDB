"""层 5：Excel 解析回归测试 —— 覆盖 Excel 链路对齐 PDF 路线的重构

本次重构：
  1. ExcelParser 简化为通用解析（去掉工程定额专用结构化逻辑）
  2. excel_to_text_stream 按 sheet 分批 yield（对齐 PDF 每页一批）
  3. 多 sheet 支持（不只读 active）
  4. sheet_start/sheet_limit 参数
"""
import sys; sys.path.insert(0, ".")
import os
import tempfile


def _create_test_excel(path: str):
    """创建测试 Excel 文件（3 个 sheet：定额表、材料表、空表）"""
    import openpyxl
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "定额表"
    ws1.append(["定额编号", "项目名称", "单位", "单价", "数量"])
    ws1.append(["A1-1", "混凝土", "m3", "350.5", "10"])
    ws1.append(["A1-2", "钢筋", "t", "4500", "2.5"])
    ws2 = wb.create_sheet("材料表")
    ws2.append(["材料编码", "材料名称", "规格", "单位"])
    ws2.append(["M001", "水泥", "P.O 42.5", "t"])
    ws2.append(["M002", "钢筋", "HRB400", "t"])
    wb.create_sheet("空表")  # 空表，验证不报错
    wb.save(path)


# ═══════════════════════════════════════════════════════════════
# 1. ExcelParser 多 sheet 解析
# ═══════════════════════════════════════════════════════════════

def test_excel_parser_multi_sheet():
    """测试 ExcelParser 遍历所有 sheet（不只读 active）"""
    from core.parser.excel_parser import ExcelParser
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        path = f.name
    try:
        _create_test_excel(path)
        doc = ExcelParser().parse(path)
        # 应该读到 3 个 sheet
        assert doc.metadata["total_sheets"] == 3, f"Expected 3 sheets, got {doc.metadata['total_sheets']}"
        sheet_names = [s["name"] for s in doc.metadata["sheets"]]
        assert sheet_names == ["定额表", "材料表", "空表"], f"Unexpected sheets: {sheet_names}"
        # raw_text 应包含所有 sheet 的内容
        assert "定额表" in doc.raw_text
        assert "材料表" in doc.raw_text
        assert "A1-1" in doc.raw_text
        assert "M001" in doc.raw_text
    finally:
        os.unlink(path)
    print("OK - ExcelParser 多 sheet 解析正确")


def test_excel_parser_no_hardcoded_logic():
    """测试 ExcelParser 不包含工程定额专用逻辑（输出是通用文本）"""
    from core.parser.excel_parser import ExcelParser
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        path = f.name
    try:
        _create_test_excel(path)
        doc = ExcelParser().parse(path)
        # structured_tables 应该为空（不预做结构化拆分，交给 AI FC）
        assert doc.structured_tables == [], f"Expected empty structured_tables, got {doc.structured_tables}"
        # tables 也应该为空
        assert doc.tables == [], f"Expected empty tables, got {doc.tables}"
        # raw_text 应该是纯文本（行号 + | 分隔）
        assert "行1:" in doc.raw_text
        assert " | " in doc.raw_text
    finally:
        os.unlink(path)
    print("OK - ExcelParser 无硬编码工程逻辑（纯通用解析）")


# ═══════════════════════════════════════════════════════════════
# 2. excel_to_text_stream 分批 yield
# ═══════════════════════════════════════════════════════════════

def test_excel_stream_all_sheets():
    """测试 excel_to_text_stream 按 sheet 分批 yield"""
    from pipeline.runner import excel_to_text_stream
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        path = f.name
    try:
        _create_test_excel(path)
        sheets = list(excel_to_text_stream(path))
        # 应该 yield 3 次（3 个 sheet）
        assert len(sheets) == 3, f"Expected 3 yields, got {len(sheets)}"
        # sheet_num 应该是 1-indexed
        assert sheets[0][0] == 1, f"Expected sheet_num=1, got {sheets[0][0]}"
        assert sheets[1][0] == 2, f"Expected sheet_num=2, got {sheets[1][0]}"
        assert sheets[2][0] == 3, f"Expected sheet_num=3, got {sheets[2][0]}"
        # 第一个 sheet 应包含定额表数据
        assert "A1-1" in sheets[0][1]
        # 第二个 sheet 应包含材料表数据
        assert "M001" in sheets[1][1]
    finally:
        os.unlink(path)
    print("OK - excel_to_text_stream 按 sheet 分批 yield（3 个 sheet）")


def test_excel_stream_sheet_range():
    """测试 excel_to_text_stream 的 sheet_start/sheet_limit 参数"""
    from pipeline.runner import excel_to_text_stream
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        path = f.name
    try:
        _create_test_excel(path)
        # 只读第 2 个 sheet（sheet_start=1, sheet_limit=1）
        sheets = list(excel_to_text_stream(path, sheet_start=1, sheet_limit=1))
        assert len(sheets) == 1, f"Expected 1 yield, got {len(sheets)}"
        assert sheets[0][0] == 2, f"Expected sheet_num=2, got {sheets[0][0]}"
        assert "M001" in sheets[0][1], "Expected 材料表 content"

        # 读前 2 个 sheet（sheet_start=0, sheet_limit=2）
        sheets = list(excel_to_text_stream(path, sheet_start=0, sheet_limit=2))
        assert len(sheets) == 2, f"Expected 2 yields, got {len(sheets)}"
        assert sheets[0][0] == 1
        assert sheets[1][0] == 2
    finally:
        os.unlink(path)
    print("OK - excel_to_text_stream sheet_start/sheet_limit 参数正确")


def test_excel_stream_empty_sheet():
    """测试 excel_to_text_stream 处理空 sheet 不报错"""
    from pipeline.runner import excel_to_text_stream
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        path = f.name
    try:
        _create_test_excel(path)
        # 空表是第 3 个 sheet，应该 yield 一次但内容可能是 "[Sheet: 空表] (1行 x 1列)"
        sheets = list(excel_to_text_stream(path, sheet_start=2, sheet_limit=1))
        assert len(sheets) == 1, f"Expected 1 yield for empty sheet, got {len(sheets)}"
    finally:
        os.unlink(path)
    print("OK - excel_to_text_stream 空表处理正常")


# ═══════════════════════════════════════════════════════════════
# 3. process_file 对 Excel 的页码提示
# ═══════════════════════════════════════════════════════════════

def test_process_file_excel_unit_display():
    """测试 process_file 对 Excel 文件显示 '块'（流单元）而非 '页'"""
    from unittest.mock import patch, MagicMock
    from agent import tools

    captured = {}
    def mock_run(filepath, industry, page_start, page_limit, batch_size, overwrite, only_tables=None, fields=None):
        captured['page_start'] = page_start
        captured['page_limit'] = page_limit
        return {"batches": [{}], "conflicts": []}

    with patch('config.settings.settings') as mock_settings, \
         patch('pipeline.runner.run', mock_run), \
         patch('pipeline.runner._check_file_size'), \
         patch('core.data_ops.get_driver') as mock_drv, \
         patch('industries.base.discover_industries'), \
         patch('industries.base.get_industry') as mock_ind:
        mock_settings.INDUSTRY = "engineering"
        mock_settings.current_file = "test.xlsx"
        mock_drv.return_value.query.return_value = [{'c': 0}]
        mock_cfg = MagicMock()
        mock_cfg.tables = [{"name": "t1"}]
        mock_ind.return_value = mock_cfg

        # Excel 文件 → 应显示 "块"
        result = tools.process_file(filepath="test.xlsx", page_start=1, page_end=3)
        assert "块" in result, f"Expected '块' in result, got {result}"
        assert "页" not in result, f"Should not contain '页' for Excel, got {result}"

        # PDF 文件 → 应显示 "页"
        result = tools.process_file(filepath="test.pdf", page_start=1, page_end=5)
        assert "页" in result, f"Expected '页' in result, got {result}"
        assert "Sheet" not in result, f"Should not contain 'Sheet' for PDF, got {result}"

    print("OK - process_file 按文件类型显示页/块")


# ═══════════════════════════════════════════════════════════════
# 4. batch_process overlap 上下文接线（跨页表格防切断）
# ═══════════════════════════════════════════════════════════════

class _OverlapFakeDriver:
    """batch_process 建 FC schema 用的最小驱动：表存在且为空（=提取表）"""
    def table_exists(self, name): return True
    def query(self, sql): return [{"c": 0}]


def _make_overlap_cfg():
    from types import SimpleNamespace
    return SimpleNamespace(
        custom_prompts={"extraction_prompt": "提取 quota_header 表数据"},
        tables=[{
            "name": "quota_header",
            "business_name": "定额主表",
            "columns": [
                {"name": "id", "type": "INTEGER"},
                {"name": "quota_id", "type": "VARCHAR", "not_null": True},
                {"name": "project_name", "type": "VARCHAR"},
            ],
            "indexes": [{"unique": True, "columns": ["quota_id"]}],
        }],
    )


# 模拟一份跨页材料明细表：第2页表头+一行，第3页是无表头的续表数据
_OVERLAP_PAGES = [
    "封面：某工程定额文件",
    "表头：quota_id | project_name\nROW:A1-1 混凝土",
    "ROW:A1-2 钢筋\nROW:A1-3 模板",   # 跨页表格续页（无表头）
    "另一个表\nROW:B2-1 水泥",
    "ROW:B2-2 砂石",
]


class _GoodAI:
    """遵守指令的 AI：只从【当前批次】段提取 ROW: 行"""
    def __init__(self):
        self.prompts = []
    def call_function(self, functions, prompt_text, system_prompt=None, max_tokens=None):
        self.prompts.append(prompt_text)
        current = prompt_text.split("【当前批次｜请提取】")[-1]
        rows = []
        for line in current.splitlines():
            line = line.strip()
            if line.startswith("ROW:"):
                qid, name = line[4:].split(None, 1)
                rows.append({"quota_id": qid, "project_name": name})
        return "output_data", {"quota_header": rows}


def _run_batch_process(ai, pages, page_limit, overlap, output_dir=None):
    from unittest.mock import patch
    from pipeline.runner import batch_process
    cfg = _make_overlap_cfg()
    stream = iter([(i + 1, t, 0) for i, t in enumerate(pages)])
    with patch("core.data_ops.get_driver", return_value=_OverlapFakeDriver()):
        return list(batch_process("fake.pdf", cfg, stream, ai,
                                  page_limit, overlap, output_dir=output_dir, route=False))


def test_batch_process_overlap_context():
    """overlap=1：上一批末页作为上下文附进 prompt 并标注勿重复提取；提取不重不漏"""
    ai = _GoodAI()
    with tempfile.TemporaryDirectory() as out_dir:
        results = _run_batch_process(ai, _OVERLAP_PAGES, page_limit=2, overlap=1,
                                     output_dir=out_dir)
        assert len(results) == 3, f"Expected 3 batches, got {len(results)}"

        # 第 1 批无上下文；第 2、3 批带上一批末页上下文
        assert "【上下文参考｜请勿提取】" not in ai.prompts[0]
        for i in (1, 2):
            assert "【上下文参考｜请勿提取】" in ai.prompts[i], f"batch{i+1} 缺上下文标注"
            assert "【当前批次｜请提取】" in ai.prompts[i]
        # 第 2 批上下文 = 第 1 批末页（第2页：表头+A1-1），当前批次含跨页续表（第3页）
        ctx2 = ai.prompts[1].split("【当前批次｜请提取】")[0]
        cur2 = ai.prompts[1].split("【当前批次｜请提取】")[1]
        assert "ROW:A1-1" in ctx2 and "表头" in ctx2, "上下文应含上一批末页(表头)"
        assert "ROW:A1-2" in cur2 and "ROW:A1-3" in cur2
        assert "ROW:A1-1" not in cur2, "上一批末页不应混入当前批次段"

        # raw 文件只存当前批次页面（不含上下文参考）
        raw2 = open(f"{out_dir}/batch_002_raw.txt", encoding="utf-8").read()
        assert "【上下文参考" not in raw2 and "ROW:A1-1" not in raw2
        assert "ROW:A1-2" in raw2

        # 跨页数据完整提取且没有因重叠重复：5 行各出现且只出现一次
        ids = [r["quota_id"] for _, d in results
               for t in d["tables"] for r in t["rows"]]
        assert sorted(ids) == ["A1-1", "A1-2", "A1-3", "B2-1", "B2-2"], ids
    print("OK - batch_process overlap=1 上下文接线（跨页表格完整提取、无重复）")


def test_batch_process_overlap_zero_unchanged():
    """overlap=0：prompt 与旧行为一致（无上下文段）"""
    ai = _GoodAI()
    results = _run_batch_process(ai, _OVERLAP_PAGES, page_limit=2, overlap=0)
    assert len(results) == 3
    for p in ai.prompts:
        assert "【上下文参考" not in p and "【当前批次" not in p
    # 第 2 批 prompt 就是第 3、4 页原文拼接
    assert ai.prompts[1] == _OVERLAP_PAGES[2] + "\n\n---\n\n" + _OVERLAP_PAGES[3]
    print("OK - batch_process overlap=0 行为不变")


def test_overlap_duplicate_extract_conflict_fallback():
    """兜底：AI 未遵守指令重复提取上下文行时，唯一键冲突检测拦截重复入库

    机制：run() 按 quota_id 分组 → insert_rows → SqliteDriver.insert
    发现 quota_id 已存在 → conflict=True → 整组跳过入库（runner.py run()）。
    注意：该拦截以行业 YAML 声明 quota_id 唯一为前提
    （sqlite_driver._get_unique_key_column 读 industries/<行业>/schemas/<表>.yaml 的唯一键声明），
    这里用 mock 模拟"已正确声明唯一约束"的行业配置。
    """
    class _BadAI(_GoodAI):
        """不遵守指令的 AI：从整个 prompt（含上下文）提取 → 跨批重复"""
        def call_function(self, functions, prompt_text, system_prompt=None, max_tokens=None):
            self.prompts.append(prompt_text)
            rows = []
            for line in prompt_text.splitlines():
                line = line.strip()
                if line.startswith("ROW:"):
                    qid, name = line[4:].split(None, 1)
                    rows.append({"quota_id": qid, "project_name": name})
            return "output_data", {"quota_header": rows}

    ai = _BadAI()
    results = _run_batch_process(ai, _OVERLAP_PAGES, page_limit=2, overlap=1)
    ids = [r["quota_id"] for _, d in results for t in d["tables"] for r in t["rows"]]
    assert len(ids) != len(set(ids)), "预期 BadAI 会产生重复提取"

    # 模拟 run() 的入库路径：临时 sqlite 库 + 声明 quota_id 唯一约束
    from unittest.mock import patch
    from core.drivers.sqlite_driver import SqliteDriver
    with tempfile.TemporaryDirectory() as tmp:
        drv = SqliteDriver(f"{tmp}/t.db")
        try:
            drv.conn.execute(
                "CREATE TABLE quota_header (id INTEGER PRIMARY KEY AUTOINCREMENT,"
                " quota_id VARCHAR, project_name VARCHAR)")
            with patch.object(SqliteDriver, "_get_unique_key_column", return_value="quota_id"):
                inserted = 0
                conflicts = 0
                for _, d in results:
                    rows = [r for t in d["tables"] for r in t["rows"]]
                    # run() 按 quota_id 分组后整组 insert，这里逐行模拟同语义
                    for r in rows:
                        res = drv.insert("quota_header", [r], overwrite=False)
                        if res.get("conflict"):
                            conflicts += 1
                        elif res.get("ok"):
                            inserted += 1
                assert inserted == 5, f"首次入库应为 5 条, got {inserted}"
                assert conflicts == len(ids) - 5, f"重复行应全部被 conflict 拦截, got {conflicts}"
                cnt = drv.query("SELECT COUNT(*) as c FROM quota_header")[0]["c"]
                assert cnt == 5, f"库里应只有 5 条（无重复入库）, got {cnt}"
        finally:
            for c in drv._conns.values():
                c.close()
    print("OK - 重复提取被唯一键冲突检测兜住（不重复入库）")


# ═══════════════════════════════════════════════════════════════
# 5. 流单元契约：Excel 500 行分块 / Word 表格 tc / 判定3 剥离语义
# ═══════════════════════════════════════════════════════════════

def test_excel_stream_row_chunking():
    """大 sheet 按 TIER1_EXCEL_ROWS 切块：跨块携带表头、行区间正确、tc=1"""
    from pipeline.runner import excel_to_text_stream
    from pipeline.constants import TIER1_EXCEL_ROWS as R
    import openpyxl
    with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
        path = f.name
    try:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "大表"
        ws.append(["编号", "名称"])
        for r in range(1, 2 * R + 2):  # 表头 + 2R+1 行数据 → 共 2R+2 行 → 3 块
            ws.append([f"Q{r:05d}", f"项目{r}"])
        wb.save(path)

        units = list(excel_to_text_stream(path))
        assert len(units) == 3, f"2R+2 行应切 3 块, got {len(units)}"
        assert all(u[2] == 1 for u in units), f"Excel 流单元 tc 应全为 1, got {[u[2] for u in units]}"
        assert [u[0] for u in units] == [1, 2, 3]
        assert "【表头参考｜请勿重复提取】" not in units[0][1]
        assert "【表头参考｜请勿重复提取】" in units[1][1]
        assert "【表头参考｜请勿重复提取】" in units[2][1]
        assert "编号" in units[1][1] and "编号" in units[2][1], "跨块应携带表头内容"
        assert f"行{R + 1}:" in units[1][1], f"第 2 块应含行{R + 1}"
        assert f"行{R + 1}:" not in units[0][1], "第 1 块不应含第 2 块的行"
        one = list(excel_to_text_stream(path, sheet_start=1, sheet_limit=1))
        assert len(one) == 1 and one[0][0] == 2
    finally:
        os.unlink(path)
    print("OK - Excel 大 sheet 按 500 行切块+跨块表头携带")


def test_docx_table_unit_tc():
    """Word 表格批 tc=1（走 AI 判定），段落批 tc=0"""
    from pipeline.runner import docx_to_text_stream
    from docx import Document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    try:
        doc = Document()
        doc.add_paragraph("第一章 说明文字")
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "编号"
        t.cell(0, 1).text = "名称"
        t.cell(1, 0).text = "A1-1"
        t.cell(1, 1).text = "混凝土"
        doc.save(path)

        units = list(docx_to_text_stream(path))
        tcs = [u[2] for u in units]
        assert 1 in tcs, f"应有表格单元 tc=1, got {tcs}"
        assert 0 in tcs, f"应有段落单元 tc=0, got {tcs}"
        table_units = [u for u in units if u[2] == 1]
        assert any("A1-1" in u[1] for u in table_units), "表格单元应含表格内容"
    finally:
        os.unlink(path)
    print("OK - Word 表格批 tc=1、段落批 tc=0")


def test_strip_table_page_lines_semantics():
    """判定3 语义：已提取业务表行被剥离，说明性零散表格行原文保留进向量库"""
    from pipeline.runner import _collect_extracted_values, _strip_table_page_lines

    batches = [{"data": {"tables": [{
        "name": "quota_items",
        "rows": [{"quota_id": "A1-1", "name": "混凝土"}]}]}}]
    values = _collect_extracted_values(batches)
    assert "A1-1" in values and "混凝土" in values

    page_text = ("[Sheet: 定额] 说明" + chr(10) +
                 "行2: A1-1 | 混凝土 | m3" + chr(10) +
                 "行3: 本章说明：定额适用于一般工业与民用建筑" + chr(10) +
                 "零散示例表: 示例编号 | 示例名称")
    stripped, total = _strip_table_page_lines([(2, page_text)], values)
    assert total == 1, f"只应剥离含已提取值的 1 行, got {total}"
    assert len(stripped) == 1
    kept = stripped[0][1]
    assert "A1-1" not in kept, "业务表行应被剥离"
    assert "本章说明" in kept, "说明文字应保留"
    assert "示例编号" in kept, "AI 未提取的零散表格行应原文保留（判定3）"

    all_table = "行1: A1-1 | 混凝土"
    stripped2, _ = _strip_table_page_lines([(3, all_table)], values)
    assert stripped2 == [], f"全剔除的单元不应进向量库, got {stripped2}"
    print("OK - 判定3：业务表行剥离、说明性表格行保留进向量库")


def test_llm_budget_guard():
    """LLM 预算护栏——超限中止、已处理批次保留、units 定位续入点"""
    from unittest.mock import patch
    from pipeline.runner import batch_process
    from pipeline.constants import TIER2_BATCH_UNITS, TIER2_OVERLAP_UNITS

    cfg = _make_overlap_cfg()
    ai = _GoodAI()
    # 5 页 → 批1: 3页(1次调用) → 批2: 2页(预算满，中止)
    pages = [(i + 1, t, 0) for i, t in enumerate(_OVERLAP_PAGES)]
    budget = {"calls": 0, "units": 0, "stopped": False, "max_calls": 1}
    with patch('core.data_ops.get_driver', return_value=_OverlapFakeDriver()):
        results = list(batch_process("fake.pdf", cfg, iter(pages), ai,
                                     TIER2_BATCH_UNITS, TIER2_OVERLAP_UNITS,
                                     route=False, budget=budget))
    assert budget["stopped"] is True, "超预算应置 stopped"
    assert budget["calls"] == 1, f"只应调用 1 次 LLM, got {budget['calls']}"
    assert budget["units"] == TIER2_BATCH_UNITS,         f"units 只计处理完成的批（3 页）, got {budget['units']}"
    assert len(results) == 1, f"只应产出 1 个批次, got {len(results)}"

    # 默认预算（不传）正常工作
    ai2 = _GoodAI()
    with patch('core.data_ops.get_driver', return_value=_OverlapFakeDriver()):
        results2 = list(batch_process("fake.pdf", cfg, iter(pages), ai2,
                                      TIER2_BATCH_UNITS, TIER2_OVERLAP_UNITS,
                                      route=False))
    assert len(results2) == 2, f"默认预算应跑完全部 2 批, got {len(results2)}"
    print("OK - LLM 预算护栏：超限中止+已处理保留+续入定位")


def test_systemic_error_guard():
    """系统性错误护栏：全部组同一原因失败 → systemic_error，不当行坏数据"""
    from pipeline.ingestion import write_batch_groups

    class _BadDrv:
        def begin(self, name): raise RuntimeError("name '_find_fk_to_main' is not defined")
        def rollback(self, name): pass
        def commit(self): pass
        def query(self, sql): return []

    data = {"tables": [{"name": "books", "rows": [
        {"book_code": "B1", "title": "t1"},
        {"book_code": "B2", "title": "t2"}]}]}
    all_results = {"failures": [], "conflicts": []}
    write_batch_groups(data, None, _BadDrv(), "books", "book_code", False, all_results)
    assert "systemic_error" in all_results,         f"全组同一原因失败应标 systemic_error: {all_results}"
    assert "_find_fk_to_main" in all_results["systemic_error"]["reason"]
    assert all_results["systemic_error"]["groups"] == 2
    print("OK - 系统性错误护栏：全组同因失败标 systemic_error")


def test_extract_completeness_retry():
    """提取行数显著低于文本估计时自动重提一次取多者"""
    from unittest.mock import patch
    from pipeline.runner import batch_process, _estimate_data_rows, _count_extracted_rows

    # 行估计启发式：数字行计入、叙述/标记行不计
    _t1 = chr(10).join(["行1: A1-1 混凝土 350.5", "本章说明：共五条", "【表头参考｜请勿重复提取】行0: x 1"])
    assert _estimate_data_rows([_t1]) == 1
    assert _count_extracted_rows({"t1": [{"a": 1}, {"a": 2}], "t2": "not-list"}) == 2

    class _FlakyAI:
        """第一次只提 1 行，重提返回 4 行"""
        def __init__(self): self.calls = 0
        def call_function(self, functions, prompt_text, system_prompt=None, max_tokens=None):
            self.calls += 1
            if self.calls == 1:
                return "output_data", {"quota_header": [{"quota_id": "A1-1", "project_name": "混凝土"}]}
            return "output_data", {"quota_header": [
                {"quota_id": f"A1-{i}", "project_name": f"项目{i}"} for i in range(1, 5)]}

    cfg = _make_overlap_cfg()
    ai = _FlakyAI()
    _p1 = chr(10).join(["行1: A1-1 混凝土 350.5", "行2: A1-2 钢筋 4500", "行3: A1-3 模板 88.2"])
    _p2 = chr(10).join(["行4: A1-4 砌块 66.1", "行5: A1-5 砂浆 12.5", "行6: A1-6 碎石 45.0"])
    pages = [(1, _p1, 1), (2, _p2, 1)]
    with patch('core.data_ops.get_driver', return_value=_OverlapFakeDriver()):
        results = list(batch_process("fake.pdf", cfg, iter(pages), ai, 2, 0, route=False))
    assert ai.calls == 2, f"行数不足应自动重提一次, 实际调用 {ai.calls} 次"
    rows = results[0][1]["tables"][0]["rows"]
    assert len(rows) == 4, f"应采用重提的更多结果（4 行）, got {len(rows)}"

    # 首次提取已足够时不重提
    class _GoodOnce:
        def __init__(self): self.calls = 0
        def call_function(self, functions, prompt_text, system_prompt=None, max_tokens=None):
            self.calls += 1
            return "output_data", {"quota_header": [
                {"quota_id": f"A1-{i}", "project_name": f"项目{i}"} for i in range(1, 6)]}
    ai2 = _GoodOnce()
    with patch('core.data_ops.get_driver', return_value=_OverlapFakeDriver()):
        list(batch_process("fake.pdf", cfg, iter(pages), ai2, 2, 0, route=False))
    assert ai2.calls == 1, f"提取充足不应重提, 实际调用 {ai2.calls} 次"
    print("OK - 提取完整性校验：低提取自动重提取多者、充足不重提")


# ═══════════════════════════════════════════════════════════════
# 主入口
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    test_excel_parser_multi_sheet()
    test_excel_parser_no_hardcoded_logic()
    test_excel_stream_all_sheets()
    test_excel_stream_sheet_range()
    test_excel_stream_empty_sheet()
    test_process_file_excel_unit_display()
    test_batch_process_overlap_context()
    test_batch_process_overlap_zero_unchanged()
    test_overlap_duplicate_extract_conflict_fallback()
    test_excel_stream_row_chunking()
    test_docx_table_unit_tc()
    test_strip_table_page_lines_semantics()
    test_llm_budget_guard()
    test_systemic_error_guard()
    test_extract_completeness_retry()
    print("\n=== ALL EXCEL REGRESSION TESTS PASSED ===")
